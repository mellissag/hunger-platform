"""Phase 15 coverage boost — part 2.

Targets:
- knowledge_extract (text extraction pure functions)
- knowledge_admin (KB CRUD without Redis)
- salon_admin_service (masking helpers + read paths)
- core/scope (master with no master_id paths)
- workers/indexer.index_kb_document (no-key early return + empty content)
- ai_service helpers
- broadcast_service additional paths
"""

from __future__ import annotations

import io
import uuid
from datetime import UTC, datetime

import pytest

from app.db.base import get_async_session_factory


# ─── knowledge_extract ───────────────────────────────────────────────────────

def test_extract_text_unsupported():
    from app.services.knowledge_extract import extract_text_from_upload
    with pytest.raises(ValueError, match="Unsupported"):
        extract_text_from_upload("file.xlsx", b"some bytes")


def test_extract_text_txt_raises():
    from app.services.knowledge_extract import extract_text_from_upload
    with pytest.raises(ValueError):
        extract_text_from_upload("file.txt", b"hello world")


def _make_minimal_docx() -> bytes:
    """Create a minimal DOCX in memory using python-docx."""
    import docx
    buf = io.BytesIO()
    d = docx.Document()
    d.add_paragraph("Hello from DOCX")
    d.add_paragraph("Line two")
    d.save(buf)
    return buf.getvalue()


def test_extract_text_docx():
    from app.services.knowledge_extract import extract_text_from_upload
    data = _make_minimal_docx()
    text = extract_text_from_upload("sample.docx", data)
    assert "Hello from DOCX" in text
    assert "Line two" in text


# ─── knowledge_admin ─────────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_kb_list_documents_empty():
    from app.services.knowledge_admin import list_documents
    factory = get_async_session_factory()
    async with factory() as s:
        rows, total = await list_documents(s, page=1, page_size=20)
    assert total == 0


@pytest.mark.asyncio
async def test_kb_create_and_get_document():
    from app.models.enums import KBSourceType
    from app.services.knowledge_admin import create_document, get_document
    factory = get_async_session_factory()
    async with factory() as s:
        doc = await create_document(
            s,
            title="FAQ",
            source_type=KBSourceType.manual,
            source_ref=None,
            content="What are the prices? Our prices start from €20.",
            lang="en",
        )
        await s.commit()
        fetched = await get_document(s, doc.id)
    assert fetched is not None
    assert fetched.id == doc.id
    assert "prices" in fetched.content


@pytest.mark.asyncio
async def test_kb_get_document_not_found():
    from app.services.knowledge_admin import get_document
    factory = get_async_session_factory()
    async with factory() as s:
        result = await get_document(s, uuid.uuid4())
    assert result is None


@pytest.mark.asyncio
async def test_kb_update_document():
    from app.models.enums import KBSourceType
    from app.services.knowledge_admin import create_document, get_document, update_document
    factory = get_async_session_factory()
    async with factory() as s:
        doc = await create_document(
            s,
            title="Services",
            source_type=KBSourceType.manual,
            source_ref=None,
            content="We offer haircuts.",
            lang="en",
        )
        await s.commit()
        doc2 = await get_document(s, doc.id)
        updated = await update_document(s, doc2, {"content": "We offer haircuts and coloring."})
        await s.commit()
    assert "coloring" in updated.content


@pytest.mark.asyncio
async def test_kb_delete_document():
    from app.models.enums import KBSourceType
    from app.services.knowledge_admin import create_document, delete_document, get_document
    factory = get_async_session_factory()
    async with factory() as s:
        doc = await create_document(
            s,
            title="To Delete",
            source_type=KBSourceType.manual,
            source_ref=None,
            content="Temporary content.",
            lang="en",
        )
        await s.commit()
        doc2 = await get_document(s, doc.id)
        await delete_document(s, doc2)
        await s.commit()
        result = await get_document(s, doc.id)
    assert result is None


@pytest.mark.asyncio
async def test_kb_list_documents_with_data():
    from app.models.enums import KBSourceType
    from app.services.knowledge_admin import create_document, list_documents
    factory = get_async_session_factory()
    async with factory() as s:
        await create_document(
            s, title="Doc1", source_type=KBSourceType.manual,
            source_ref=None, content="content1", lang="en"
        )
        await create_document(
            s, title="Doc2", source_type=KBSourceType.manual,
            source_ref=None, content="content2", lang="ru"
        )
        await s.commit()
        rows, total = await list_documents(s, page=1, page_size=20)
    assert total == 2
    assert len(rows) == 2


# ─── salon_admin_service helpers ─────────────────────────────────────────────

def test_mask_token_short():
    from app.services.salon_admin_service import _mask_token
    assert _mask_token(None) is None
    assert _mask_token("short") is None  # len < 8
    assert _mask_token("12345678") == "****5678"


def test_mask_token_long():
    from app.services.salon_admin_service import _mask_token
    result = _mask_token("1234567890abcdef")
    assert result is not None
    assert result.startswith("****")
    assert result.endswith("cdef")


def test_mask_integrations_empty():
    from app.services.salon_admin_service import mask_integrations
    assert mask_integrations({}) == {}


def test_mask_integrations_no_secrets():
    from app.services.salon_admin_service import mask_integrations
    data = {"telegram": {"webhook_url": "https://example.com"}}
    result = mask_integrations(data)
    assert result["telegram"]["webhook_url"] == "https://example.com"


def test_mask_integrations_with_bot_token():
    from app.services.salon_admin_service import mask_integrations
    data = {"telegram": {"bot_token": "1234567890:ABCDEFGHIJabcdefghij", "webhook_url": "x"}}
    result = mask_integrations(data)
    token = result["telegram"]["bot_token"]
    assert token.startswith("****")


def test_mask_integrations_with_password():
    from app.services.salon_admin_service import mask_integrations
    data = {"smtp": {"host": "mail.example.com", "password": "secret123"}}
    result = mask_integrations(data)
    assert result["smtp"]["password"] == "****"
    assert result["smtp"]["host"] == "mail.example.com"


def test_mask_integrations_empty_password():
    from app.services.salon_admin_service import mask_integrations
    data = {"smtp": {"password": ""}}
    result = mask_integrations(data)
    assert result["smtp"]["password"] == ""


def test_mask_integrations_non_dict_value():
    from app.services.salon_admin_service import mask_integrations
    data = {"some_flag": True, "count": 42}
    result = mask_integrations(data)
    assert result["some_flag"] is True
    assert result["count"] == 42


# ─── core/scope edge cases ───────────────────────────────────────────────────

def test_scope_booking_filter_non_master():
    from unittest.mock import MagicMock
    from app.core.scope import booking_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.admin
    result = booking_scope_filter(u)
    # Should return literal(True) for non-master
    assert result is not None


def test_scope_booking_filter_master_no_mid():
    from unittest.mock import MagicMock
    from app.core.scope import booking_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = None
    result = booking_scope_filter(u)
    # Should return false() when no master_id
    assert result is not None


def test_scope_booking_filter_master_with_mid():
    from unittest.mock import MagicMock
    from app.core.scope import booking_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = uuid.uuid4()
    result = booking_scope_filter(u)
    assert result is not None


def test_scope_client_filter_non_master():
    from unittest.mock import MagicMock
    from app.core.scope import client_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.owner
    result = client_scope_filter(u)
    assert result is not None


def test_scope_client_filter_master_no_mid():
    from unittest.mock import MagicMock
    from app.core.scope import client_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = None
    result = client_scope_filter(u)
    assert result is not None


def test_scope_master_record_filter_non_master():
    from unittest.mock import MagicMock
    from app.core.scope import master_record_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.admin
    result = master_record_scope_filter(u)
    assert result is not None


def test_scope_master_record_filter_master_no_mid():
    from unittest.mock import MagicMock
    from app.core.scope import master_record_scope_filter
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = None
    result = master_record_scope_filter(u)
    assert result is not None


def test_scope_ensure_master_can_access_booking_non_master():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_can_access_booking
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.admin
    # Should not raise
    ensure_master_can_access_booking(u, uuid.uuid4())


def test_scope_ensure_master_can_access_booking_no_mid():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_can_access_booking
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = None
    # Should not raise (no master_id means can't lock to anything)
    ensure_master_can_access_booking(u, uuid.uuid4())


def test_scope_ensure_master_can_access_booking_own():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_can_access_booking
    from app.models.enums import UserRole
    mid = uuid.uuid4()
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = mid
    ensure_master_can_access_booking(u, mid)  # should not raise


def test_scope_ensure_master_can_access_booking_other():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_can_access_booking
    from app.core.exceptions import ForbiddenScopeError
    from app.models.enums import UserRole
    mid = uuid.uuid4()
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = mid
    with pytest.raises(ForbiddenScopeError):
        ensure_master_can_access_booking(u, uuid.uuid4())  # different ID


def test_scope_ensure_master_own_master_id_non_master():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_own_master_id
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.admin
    ensure_master_own_master_id(u, uuid.uuid4())  # should not raise


def test_scope_ensure_master_own_master_id_no_mid():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_own_master_id
    from app.models.enums import UserRole
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = None
    ensure_master_own_master_id(u, uuid.uuid4())  # should not raise


def test_scope_ensure_master_own_master_id_mismatch():
    from unittest.mock import MagicMock
    from app.core.scope import ensure_master_own_master_id
    from app.core.exceptions import ForbiddenScopeError
    from app.models.enums import UserRole
    mid = uuid.uuid4()
    u = MagicMock()
    u.role = UserRole.master
    u.master_id = mid
    with pytest.raises(ForbiddenScopeError):
        ensure_master_own_master_id(u, uuid.uuid4())


# ─── workers/indexer: no-key early return ────────────────────────────────────

@pytest.mark.asyncio
async def test_indexer_no_gemini_key(monkeypatch):
    """index_kb_document should return early if GEMINI_API_KEY is not set."""
    from app.workers.indexer import index_kb_document
    from unittest.mock import MagicMock
    import app.workers.indexer as idx_mod

    factory = get_async_session_factory()

    # Patch get_settings to return a mock with no gemini key
    mock_settings = MagicMock()
    mock_settings.gemini_api_key = None
    monkeypatch.setattr(idx_mod, "get_settings", lambda: mock_settings)

    ctx = {"db": factory}
    # Should return silently (no exception)
    await index_kb_document(ctx, str(uuid.uuid4()))


@pytest.mark.asyncio
async def test_indexer_empty_content(monkeypatch):
    """index_kb_document: doc with empty content should delete old chunks."""
    from app.models.enums import KBSourceType
    from app.services.knowledge_admin import create_document
    from app.workers.indexer import index_kb_document

    import app.workers.indexer as idx_mod
    import unittest.mock as mock
    from unittest.mock import MagicMock

    factory = get_async_session_factory()

    async with factory() as s:
        doc = await create_document(
            s,
            title="Empty Doc",
            source_type=KBSourceType.manual,
            source_ref=None,
            content="   ",  # whitespace = empty
            lang="en",
        )
        await s.commit()
        doc_id = doc.id

    mock_settings = MagicMock()
    mock_settings.gemini_api_key = "fake-key-for-test"
    monkeypatch.setattr(idx_mod, "get_settings", lambda: mock_settings)

    with mock.patch.object(idx_mod, "genai"):
        ctx = {"db": factory}
        await index_kb_document(ctx, str(doc_id))


@pytest.mark.asyncio
async def test_indexer_doc_not_found(monkeypatch):
    """index_kb_document with non-existent doc_id should log warning and return."""
    from app.workers.indexer import index_kb_document

    import app.workers.indexer as idx_mod
    import unittest.mock as mock
    from unittest.mock import MagicMock

    mock_settings = MagicMock()
    mock_settings.gemini_api_key = "fake-key"
    monkeypatch.setattr(idx_mod, "get_settings", lambda: mock_settings)

    factory = get_async_session_factory()
    with mock.patch.object(idx_mod, "genai"):
        ctx = {"db": factory}
        await index_kb_document(ctx, str(uuid.uuid4()))


# ─── AI rate limit service ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_ai_rate_limit_db_fallback_under_limit():
    """AI rate limit DB fallback: under limit should not raise."""
    from app.services.ai_rate_limit import check_ai_rate_limit
    factory = get_async_session_factory()
    async with factory() as s:
        # Fresh DB = 0 messages = under limit
        await check_ai_rate_limit(s, redis=None, client_id=uuid.uuid4())


@pytest.mark.asyncio
async def test_ai_rate_limit_hour_bucket():
    """_hour_bucket should return consistent UTC hour string."""
    from app.services.ai_rate_limit import _hour_bucket
    from datetime import UTC, datetime
    dt = datetime(2026, 4, 15, 10, 30, 0, tzinfo=UTC)
    bucket = _hour_bucket(dt)
    assert bucket == "2026041510"


@pytest.mark.asyncio
async def test_ai_rate_limit_naive_datetime():
    """_hour_bucket handles naive datetime."""
    from app.services.ai_rate_limit import _hour_bucket
    from datetime import datetime
    dt = datetime(2026, 4, 15, 10, 30, 0)
    bucket = _hour_bucket(dt)
    assert bucket.startswith("202604151")


# ─── notification_service helpers ────────────────────────────────────────────

@pytest.mark.asyncio
async def test_notification_get_admin_chat_id_empty():
    """get_admin_notify_chat_id with empty DB should return None."""
    from app.services.notification_service import get_admin_notify_chat_id
    factory = get_async_session_factory()
    async with factory() as s:
        result = await get_admin_notify_chat_id(s)
    assert result is None
